import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import uuid
import tempfile

PROPOSAL_CONFIG = {
    "SMM, Meta & Google Ads and SEO": {
        "template": "SMM, Meta & Google Ads and SEO.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Social Media Setup", "SM", "One Time Fee"),
            ("Meta & Google Ads Setup", "MG", "One Time Fee"),
            ("Creative Posts", "CP", "Monthly Recurring Fee"),
            ("Meta Paid Ads", "MP", "Monthly Recurring Fee"),
            ("Google Paid Ads", "GP", "Monthly Recurring Fee"),
            ("SEO", "SEO", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "SEO & Google Ads Campaign": {
        "template": "SEO & Google ads campaign.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Google Ads Setup", "MG", "One Time Fee"),
            ("Google Paid Ads", "GP", "Monthly Recurring Fee"),
            ("SEO Optimization", "SEO", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "Only Google Ads Campaign": {
        "template": "Only Google Ads Campaign.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Google Ads Setup", "MG", "One Time Fee"),
            ("Google Paid Ads", "GP", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "SMM, Google ads & Meta ads Campaigns": {
        "template": "SMM, Google ads & Meta ads Campaigns.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Social Media Setup", "SM", "One Time Fee"),
            ("Meta & Google Ads Setup", "MG", "One Time Fee"),
            ("Creative Posts", "CP", "Monthly Recurring Fee"),
            ("Meta Paid Ads", "MP", "Monthly Recurring Fee"),
            ("Google Paid Ads", "GP", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "Only Email Marketing": {
        "template": "Only Email Marketing.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Email Marketing", "EM", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "Only SEO": {
        "template": "Only SEO.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("SEO", "SEO", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "SMM & Meta Ads Campaigns": {
        "template": "SMM & Meta Ads Campaigns.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Social Media Setup", "SM", "One Time Fee"),
            ("Meta & Google Ads Setup", "MG", "One Time Fee"),
            ("Creative Posts", "CP", "Monthly Recurring Fee"),
            ("Meta Paid Ads", "MP", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "Only Meta Ads Campaigns": {
        "template": "Only Meta Ads Campaigns.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Social Media Setup", "SM", "One Time Fee"),
            ("Meta Ads Setup", "MG", "One Time Fee"),
            ("Meta Paid Ads", "MP", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "Only SMM": {
        "template": "Only SMM.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Social Media Setup", "SM", "One Time Fee"),
            ("Creative Posts", "CP", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    },
    "All Services (DM Proposal - All)": {
        "template": "DM Proposal - All.docx",
        "pricing_fields": [
            ("Marketing Strategy", "MS", "One Time Fee"),
            ("Social Media Setup", "SM", "One Time Fee"),
            ("Meta & Google Ads Setup", "MG", "One Time Fee"),
            ("Creative Posts", "CP", "Monthly Recurring Fee"),
            ("Meta Paid Ads", "MP", "Monthly Recurring Fee"),
            ("Google Paid Ads", "GP", "Monthly Recurring Fee"),
            ("SEO", "SEO", "Monthly Recurring Fee"),
            ("Email Marketing", "EM", "Monthly Recurring Fee"),
            ("Monthly Reporting", "MR", "Monthly Recurring Fee")
        ],
        "team_type": "dm",
        "special_fields": [("VDate", "<<")],
        "proposal_type": "digital_marketing"
    }
}

def apply_formatting(new_run, original_run):
    """Copy formatting from original run to new run"""
    if original_run.font.name:
        new_run.font.name = original_run.font.name
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), original_run.font.name)
    if original_run.font.size:
        new_run.font.size = original_run.font.size
    if original_run.font.color.rgb:
        new_run.font.color.rgb = original_run.font.color.rgb
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic

def replace_in_paragraph(para, placeholders):
    """Handle paragraph replacements preserving formatting"""
    original_runs = para.runs.copy()
    full_text = para.text
    for ph, value in placeholders.items():
        full_text = full_text.replace(ph, str(value))

    if full_text != para.text:
        para.clear()
        new_run = para.add_run(full_text)
        if original_runs:
            original_run = next((r for r in original_runs if r.text), None)
            if original_run:
                apply_formatting(new_run, original_run)

def replace_and_format(doc, placeholders):
    """Enhanced replacement with table cell handling"""
    for para in doc.paragraphs:
        replace_in_paragraph(para, placeholders)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    replace_in_paragraph(para, placeholders)
                else:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, placeholders)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return doc

def get_dm_team_details():
    """Collect team composition for DM projects"""
    st.subheader("Team Composition")
    team_roles = {
        "Digital Marketing Executive": "DME",
        "Digital Marketing Associate": "DMA",
        "Business Analyst": "BA",
        "Graphics Designer": "GD"
    }
    team_details = {}
    cols = st.columns(2)

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def remove_empty_rows(table):
    """Remove entire rows where the last cell is empty after replacements"""
    rows_to_remove = []
    # Keep header row (assuming first row is header)
    header_row = table.rows[0] if table.rows else None
    
    for row in table.rows:
        # Skip header row
        if row == header_row:
            continue
            
        # Check last cell for emptiness
        if row.cells:
            last_cell = row.cells[-1]
            last_cell_text = last_cell.text.strip()
            if not last_cell_text:  # Cell is empty
                rows_to_remove.append(row)
    
    # Remove identified rows
    for row in reversed(rows_to_remove):
        table._tbl.remove(row._element)

def validate_phone_number(country, phone_number):
    """Validate phone number format"""
    if country.lower() == "india":
        return phone_number.startswith("+91")
    return phone_number.startswith("+1")

def format_number_with_commas(number):
    return f"{number:,}"

def calculate_installments(pricing_fields, numerical_values, currency):
    """Calculate installment amounts based on fee types"""
    one_time_fee_total = 0
    monthly_fee_total = 0
    
    # Calculate totals for each fee type
    for label, key, fee_type in pricing_fields:
        value = numerical_values.get(key, 0)
        if "One Time" in fee_type:
            one_time_fee_total += value
        elif "Monthly" in fee_type:
            monthly_fee_total += value
    
    # Calculate installments with GST for INR
    if currency == "INR":
        instalment1 = one_time_fee_total + (one_time_fee_total * 0.18)
        instalment2 = monthly_fee_total + (monthly_fee_total * 0.18)
    else:
        instalment1 = one_time_fee_total
        instalment2 = monthly_fee_total
    
    return int(instalment1), int(instalment2)

def generate_document():
    st.title("DM Proposal Generator")
    base_dir = os.getcwd()

    selected_proposal = st.selectbox("Select Proposal", list(PROPOSAL_CONFIG.keys()))
    config = PROPOSAL_CONFIG[selected_proposal]
    template_path = os.path.join(base_dir, config["template"])

    # Client Information
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Client Name:")
        client_email = st.text_input("Client Email:")
    with col2:
        client_number = st.text_input("Client Number:")
    
    date_field = st.date_input("Date:", datetime.today())

    # Currency Handling
    currency = st.selectbox("Select Currency", ["USD", "INR"])
    currency_symbol = "$" if currency == "USD" else "₹"

    # Special Fields
    special_data = {}
    st.subheader("Additional Details")
    vdate = st.date_input("Proposal Validity Until:")
    special_data["<<VDate>>"] = vdate.strftime("%d-%m-%Y")

    # Pricing Section
    st.subheader("Pricing Details")
    pricing_data = {}
    numerical_values = {}
    
    pricing_fields = config["pricing_fields"]
    num_fields = len(pricing_fields)
    num_rows = (num_fields + 1) // 2

    for row in range(num_rows):
        cols = st.columns(2)
        for col in range(2):
            idx = row * 2 + col
            if idx < num_fields:
                label, key, fee_type = pricing_fields[idx]
                with cols[col]:
                    value = st.number_input(
                        f"{label} ({fee_type})",
                        min_value=0,
                        value=0,
                        step=100,
                        format="%d",
                        key=f"price_{key}"
                    )
                    numerical_values[key] = value
                    pricing_data[f"<<{key}>>"] = f"{currency_symbol}{format_number_with_commas(value)}" if value > 0 else ""

    # Calculate totals
    services_sum = sum(numerical_values.values())
    gst = int(services_sum * 0.18) if currency == "INR" else 0
    tp = int(services_sum + gst) if currency == "INR" else services_sum

    pricing_data["<<Total>>"] = f"{currency_symbol}{format_number_with_commas(services_sum)}" if currency == "INR" else ""
    pricing_data["<<GST>>"] = f"{currency_symbol}{format_number_with_commas(gst)}" if currency == "INR" else ""
    pricing_data["<<TP>>"] = f"{currency_symbol}{format_number_with_commas(tp)}"

    # Payment Schedule Section - Now automated
    st.subheader("Payment Schedule")
    instalment1, instalment2 = calculate_installments(pricing_fields, numerical_values, currency)
    
    # Display calculated installments (read-only) with conditional GST text
    gst_text = " + GST" if currency == "INR" else ""
    st.write(f"Instalment 1 (One-time fee total{gst_text}): {currency_symbol}{format_number_with_commas(instalment1)}")
    st.write(f"Instalment 2 (Monthly recurring fee total{gst_text}): {currency_symbol}{format_number_with_commas(instalment2)}")
    
    instalment_data = {
        "<<Instalment 1>>": f"{currency_symbol}{format_number_with_commas(instalment1)}",
        "<<Instalment 2>>": f"{currency_symbol}{format_number_with_commas(instalment2)}"
    }

    # Team Composition
    team_data = get_dm_team_details()

    # Combine all placeholders
    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Email>>": client_email,
        "<<Client Number>>": client_number,
        "<<Date>>": date_field.strftime("%d-%m-%Y"),
    }
    placeholders.update(pricing_data)
    placeholders.update(team_data)
    placeholders.update(special_data)
    placeholders.update(instalment_data)

    if st.button("Generate Proposal"):
        doc_filename = f"DM Proposal - {client_name} {date_field.strftime('%d-%m-%Y')}.docx"

        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                doc = Document(template_path)
            except FileNotFoundError:
                st.error(f"Template not found: {template_path}")
                return

            doc = replace_and_format(doc, placeholders)

            # Remove GST and Total lines for USD
            if currency == "USD":
                paragraphs = list(doc.paragraphs)
                for para in reversed(paragraphs):
                    if "Total Marketing Cost" in para.text or "GST" in para.text:
                        p = para._element
                        p.getparent().remove(p)

            for table in doc.tables:
                remove_empty_rows(table)

            doc_path = os.path.join(temp_dir, doc_filename)
            doc.save(doc_path)

            with open(doc_path, "rb") as f:
                st.download_button(
                    label="Download Proposal",
                    data=f,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    generate_document()